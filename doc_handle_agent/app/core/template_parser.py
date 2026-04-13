"""模板文档解析器."""

import json
import re
from typing import Dict, List, Optional, Union

from app.core.state import StaticParagraph, TemplateParagraph
from app.infrastructure.docx_handler import DocxHandler, ParagraphInfo
from app.utils.logger import get_logger

logger = get_logger(__name__)


class TemplateParser:
    """模板文档解析器."""

    def __init__(self):
        """初始化模板解析器."""
        self.block_pattern = re.compile(r'\{\{(.*?)\}\}')
        self.docx_handler = DocxHandler()

    def parse(self, template_path: str) -> List[TemplateParagraph]:
        """解析模板文件，提取所有模板段落.

        解析逻辑：
        1. 提取所有段落信息
        2. 识别模板段落（包含 {{...}} 的段落）
        3. 顺序向下寻找低层级标题模板，找到的段落放到其上一级段落的 children 中
        4. 子段落可以是静态的或模板的

        Args:
            template_path: 模板文件路径

        Returns:
            模板段落列表
        """
        logger.info(
            "parse_template_start",
            template_path=template_path,
        )

        try:
            # 提取所有段落信息
            paragraphs_info = self.docx_handler.extract_paragraphs_info(template_path)

            # 解析模板段落
            paragraphs, _ = self._parse_paragraphs(
                paragraphs_info, start_index=0, parent_level=0
            )

            logger.info(
                "parse_template_success",
                template_path=template_path,
                paragraph_count=len(paragraphs),
            )

            return paragraphs

        except Exception as e:
            logger.error(
                "parse_template_failed",
                template_path=template_path,
                error=str(e),
            )
            raise

    def _parse_paragraphs(
        self,
        paragraphs_info: List[ParagraphInfo],
        start_index: int,
        parent_level: int,
    ) -> tuple[List[TemplateParagraph], int]:
        """解析模板段落.

        算法步骤：
        1. 判断当前段落是否为模板段落，不是则跳过
        2. 若是模板段落且为正文则直接解析
        3. 若是模板段落且为标题，则向下遍历直到找到标题等级小于或等于当前段落的段落，
           将遍历出来的段落按照标题等级构建层级结构

        Args:
            paragraphs_info: 所有段落信息
            start_index: 开始解析的索引
            parent_level: 父段落级别（用于判断子段落边界）

        Returns:
            (解析出的模板段落列表, 下一个待解析的索引)
        """
        paragraphs: List[TemplateParagraph] = []
        i = start_index

        while i < len(paragraphs_info):
            para_info = paragraphs_info[i]

            # 如果遇到同层级或更高层级的标题，停止（仅限于非顶级调用）
            if para_info.is_heading and parent_level > 0:
                current_level = self._get_heading_level(para_info.style_name)
                if current_level <= parent_level:
                    break

            # 步骤1: 判断当前段落是否为模板段落，不是则跳过
            if not para_info.has_template:
                i += 1
                continue

            # 解析模板段落
            paragraph = self._parse_template_paragraph(para_info)
            current_level = self._get_heading_level(para_info.style_name)

            # 步骤2: 若是模板段落且为正文则直接解析（作为顶层段落）
            if not para_info.is_heading:
                paragraphs.append(paragraph)
                i += 1
                continue

            # 步骤3: 若是模板段落且为标题，收集其子段落
            # 向下遍历直到找到标题等级小于或等于当前段落的段落
            children, next_index = self._collect_children_for_heading(
                paragraphs_info, i + 1, current_level
            )
            paragraph.children = children
            paragraphs.append(paragraph)
            i = next_index

        return paragraphs, i

    def _collect_children_for_heading(
        self,
        paragraphs_info: List[ParagraphInfo],
        start_index: int,
        parent_level: int,
    ) -> tuple[List[Union[TemplateParagraph, StaticParagraph]], int]:
        """为标题段落收集子段落.

        使用栈实现层级结构构建：
        1. 先把开始段落压栈
        2. 向下遍历
        3. 当段落标题等级高于栈顶段落时（数值更小，级别更高）：
           - 将该段落添加到栈顶段落的子段落中
           - 将当前段落压栈
        4. 当段落标题低于或等于栈顶段落时（数值更大或相等，级别更低或相同）：
           - 把栈顶段落弹出栈
           - 再次比较栈顶段落与当前段落
           - 直到遇到标题等级高于其的栈顶段落作为其的子段落
           - 把当前段落压栈
        5. 若段落为正文段落：
           - 直接作为栈顶段落的子段落
           - 不把当前段落压栈

        Args:
            paragraphs_info: 所有段落信息
            start_index: 开始收集的索引
            parent_level: 父段落标题级别

        Returns:
            (子段落列表, 下一个待解析的索引)
        """
        children: List[Union[TemplateParagraph, StaticParagraph]] = []
        i = start_index

        # 使用栈来跟踪当前路径上的段落
        # 每个元素: (level, paragraph)
        # 注意：level 越小表示标题等级越高（Heading 1 < Heading 2 < Heading 3...）
        level_stack: List[tuple[int, Union[TemplateParagraph, StaticParagraph]]] = []

        while i < len(paragraphs_info):
            para_info = paragraphs_info[i]
            current_level = self._get_heading_level(para_info.style_name)

            # 如果遇到标题等级 <= 父段落等级的标题，停止
            # parent_level 是父段落的等级，current_level <= parent_level 表示当前是同级或更高层级的标题
            if para_info.is_heading and current_level <= parent_level:
                break

            # 处理非模板段落（静态段落）
            if not para_info.has_template:
                static_para = StaticParagraph(
                    id=str(para_info.index),
                    content=para_info.text,
                    style_name=para_info.style_name,
                    is_heading=para_info.is_heading,
                )

                # 如果是正文段落（非标题），直接作为栈顶段落的子段落，不压栈
                if not para_info.is_heading:
                    if level_stack:
                        # 找到栈顶段落作为父段落
                        _, parent = level_stack[-1]
                        parent.children.append(static_para)
                    else:
                        # 栈为空，直接添加到 children
                        children.append(static_para)
                    i += 1
                    continue

                # 是标题静态段落，需要按层级处理
                # 步骤4: 当段落标题等级低于或等于栈顶时（数值更小或相等，即级别更高或相同），弹出栈顶
                while level_stack and current_level <= level_stack[-1][0]:
                    level_stack.pop()

                # 步骤3: 将段落添加到栈顶段落的子段落中（如果栈不为空）
                if level_stack:
                    _, parent = level_stack[-1]
                    parent.children.append(static_para)
                else:
                    children.append(static_para)

                # 将当前段落压栈
                level_stack.append((current_level, static_para))
                i += 1
                continue

            # 解析模板段落
            paragraph = self._parse_template_paragraph(para_info)

            # 如果是正文模板段落，直接作为栈顶段落的子段落，不压栈
            if not para_info.is_heading:
                if level_stack:
                    # 找到栈顶段落作为父段落
                    _, parent = level_stack[-1]
                    parent.children.append(paragraph)
                else:
                    children.append(paragraph)
                i += 1
                continue

            # 是标题模板段落
            # 步骤4: 当段落标题等级低于或等于栈顶时（数值更小或相等，即级别更高或相同），弹出栈顶
            while level_stack and current_level <= level_stack[-1][0]:
                level_stack.pop()

            # 步骤3: 将段落添加到栈顶段落的子段落中（如果栈不为空）
            if level_stack:
                _, parent = level_stack[-1]
                parent.children.append(paragraph)
            else:
                children.append(paragraph)

            # 递归收集该标题的子段落
            grandchildren, next_index = self._collect_children_for_heading(
                paragraphs_info, i + 1, current_level
            )
            paragraph.children = grandchildren

            # 将当前段落压栈
            level_stack.append((current_level, paragraph))
            i = next_index

        return children, i

    def _parse_template_paragraph(self, para_info: ParagraphInfo) -> TemplateParagraph:
        """解析单个模板段落.

        Args:
            para_info: 段落信息

        Returns:
            解析后的模板段落
        """
        # 解析模板内容
        template_content = para_info.template_content or ""
        content = template_content.strip()

        # 去掉可能存在的双花括号包裹
        # DocxHandler 已经将 {{...}} 转换为 {...}，所以 content 应该是 {...} 格式
        if content.startswith("{{") and content.endswith("}}"):
            content = content[2:-2].strip()
            # 重新添加单花括号以形成合法 JSON
            content = "{" + content + "}"

        # 规范化引号：将中文引号替换为 ASCII 引号
        content = self._normalize_quotes(content)

        # 解析JSON
        try:
            data = json.loads(content)
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid template JSON at paragraph {para_info.index}: {content}") from e

        if not isinstance(data, dict):
            raise ValueError(f"Template content must be a JSON object at paragraph {para_info.index}")

        if "prompt" not in data:
            raise ValueError(f"Template must have 'prompt' field at paragraph {para_info.index}")

        # 解析属性
        is_list = str(data.get("list", "false")).lower() == "true"
        min_length = self._parse_int_field(data.get("min_length"))
        max_length = self._parse_int_field(data.get("max_length"))
        img = data.get("img")  # 图片获取提示词
        example = data.get("example")  # 内容生成参考示例

        return TemplateParagraph(
            id=str(para_info.index),
            is_template=True,
            text=para_info.text,
            style_name=para_info.style_name,
            is_heading=para_info.is_heading,
            prompt=data["prompt"],
            is_list=is_list,
            min_length=min_length,
            max_length=max_length,
            img=img,
            example=example,
            children=[],  # 子段落将在后续步骤中填充
        )

    def _get_heading_level(self, style_name: str) -> int:
        """获取标题级别."""
        if not style_name.startswith("Heading"):
            return 99

        try:
            level = int(style_name.replace("Heading", "").strip())
            return level
        except ValueError:
            return 99

    def _parse_int_field(self, value: Optional[str]) -> Optional[int]:
        """解析整数字段."""
        if value is None:
            return None
        try:
            return int(value)
        except (ValueError, TypeError):
            return None

    def _normalize_quotes(self, content: str) -> str:
        """规范化引号：将中文引号替换为 ASCII 引号.

        Args:
            content: 原始内容

        Returns:
            规范化后的内容
        """
        # 中文双引号替换为 ASCII 双引号
        content = content.replace('“', '"').replace('”', '"')
        # 中文单引号替换为 ASCII 单引号
        content = content.replace("‘", "'").replace("’", "'")
        # 中文逗号替换为 ASCII 逗号
        content = content.replace('，', ',')
        # 中文冒号替换为 ASCII 冒号
        content = content.replace('：', ':')
        return content

    def preview_blocks(self, template_path: str) -> List[Dict]:
        """预览模板中的内容块."""
        paragraphs = self.parse(template_path)
        return [p.to_dict() for p in paragraphs]

    def validate_template(self, template_path: str) -> tuple[bool, Optional[str]]:
        """验证模板文件."""
        is_valid, error = self.docx_handler.validate_template(template_path)
        if not is_valid:
            return False, error

        try:
            paragraphs = self.parse(template_path)

            if not paragraphs:
                return True, "Warning: No template paragraphs found in template"

            for para in paragraphs:
                if not para.prompt.strip():
                    return False, f"Empty prompt in paragraph {para.id}"

            return True, None

        except Exception as e:
            return False, f"Failed to parse template: {str(e)}"


def create_example_template(output_path: str) -> str:
    """创建示例模板文件."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # 添加标题
    title = doc.add_heading("系统设计文档", level=1)

    # 1. 系统概述 - 普通正文模板
    doc.add_heading("1. 系统概述", level=2)
    doc.add_paragraph('{{"prompt":"系统的整体功能概述，包括主要业务目标和技术目标"}}')

    # 2. 功能模块 - 列表模板 + 子段落
    doc.add_heading("2. 功能模块", level=2)
    doc.add_paragraph('{{"prompt":"系统的主要功能模块", "list":"true"}}')
    # 子段落：2.1 标识（静态）+ 模板（动态）
    doc.add_paragraph('    2.1 标识 {{"prompt":"随机10位英文字母序列"}}', style='List Paragraph')
    # 子段落：2.2 概要（静态）+ 模板（动态）
    doc.add_paragraph('    2.2 概要 {{"prompt":"功能模块的功能概要"}}', style='List Paragraph')

    # 3. 核心流程 - 普通正文模板
    doc.add_heading("3. 核心流程", level=2)
    doc.add_paragraph('{{"prompt":"系统的核心业务处理流程说明"}}')

    # 4. 架构设计 - 带参考示例的模板
    doc.add_heading("4. 架构设计", level=2)
    doc.add_paragraph(
        '{{'
        '"prompt":"系统的技术架构设计，包括分层架构和组件关系", '
        '"example":"本系统采用经典的三层架构设计。表现层负责用户交互，使用Vue.js框架实现响应式界面；业务逻辑层处理核心业务规则，采用Spring Boot构建RESTful API；数据访问层负责持久化操作，使用MyBatis进行ORM映射。各层之间通过定义良好的接口进行通信，降低了耦合度，提高了系统的可维护性。"'
        '}}'
    )

    # 保存
    import os
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    logger.info(
        "example_template_created",
        output_path=output_path,
    )

    return output_path
