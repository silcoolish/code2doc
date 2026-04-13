"""docx文档处理器."""

import re
from pathlib import Path
from typing import TYPE_CHECKING, Dict, List, Optional, Tuple, Union

from docx import Document
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from app.utils.logger import get_logger

if TYPE_CHECKING:
    from app.core.state import GeneratedContentResult, ImageInfo

logger = get_logger(__name__)


class ParagraphInfo:
    """段落信息."""

    def __init__(
        self,
        index: int,
        text: str,
        style_name: str,
        is_heading: bool,
        has_template: bool,
        template_content: Optional[str] = None,
    ):
        self.index = index
        self.text = text
        self.style_name = style_name
        self.is_heading = is_heading
        self.has_template = has_template
        self.template_content = template_content


class DocxHandler:
    """docx文档处理器."""

    def __init__(self):
        """初始化docx处理器."""
        self.block_pattern = re.compile(r'\{\{(.*?)\}\}')

    def read_paragraphs(self, doc_path: str) -> List[str]:
        """读取文档所有段落文本.

        Args:
            doc_path: 文档路径

        Returns:
            段落文本列表
        """
        doc = Document(doc_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        logger.info(
            "read_paragraphs",
            doc_path=doc_path,
            paragraph_count=len(paragraphs),
        )

        return paragraphs

    def extract_blocks_with_positions(
        self,
        doc_path: str,
    ) -> List[Tuple[int, str, str, bool]]:
        """提取内容块及其位置信息和段落样式类型.

        Args:
            doc_path: 文档文件路径

        Returns:
            [(段落索引, 原始文本, 块内容, 是否为标题), ...]
            是否为标题: 根据段落样式判断（Heading开头为标题，其他为正文）
        """
        doc = Document(doc_path)
        blocks = []

        for idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue

            matches = self.block_pattern.findall(text)
            if matches:
                # 判断段落样式：Heading 开头表示标题，其他为正文
                style_name = paragraph.style.name if paragraph.style else ""
                is_heading = style_name.startswith("Heading")

                for match in matches:
                    blocks.append((idx, text, match.strip(), is_heading))

        logger.info(
            "extract_blocks",
            doc_path=doc_path,
            block_count=len(blocks),
        )

        return blocks

    def extract_paragraphs_info(self, doc_path: str) -> List[ParagraphInfo]:
        """提取所有段落信息.

        Args:
            doc_path: 文档路径

        Returns:
            段落信息列表
        """
        doc = Document(doc_path)
        paragraphs_info = []

        for idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            style_name = paragraph.style.name if paragraph.style else ""
            is_heading = style_name.startswith("Heading")

            matches = self.block_pattern.findall(text)
            has_template = len(matches) > 0
            template_content = "{" + self.replace_dot(matches[0]).strip() + "}" if matches else None

            paragraphs_info.append(ParagraphInfo(
                index=idx,
                text=text,
                style_name=style_name,
                is_heading=is_heading,
                has_template=has_template,
                template_content=template_content,
            ))

        return paragraphs_info

    def replace_blocks(
        self,
        template_path: str,
        output_path: str,
        generated_contents: Dict[str, List["GeneratedContentResult"]],
        generated_images: Optional[Dict[str, List["ImageInfo"]]] = None,
    ) -> str:
        """替换模板中的内容块.

        Args:
            template_path: 模板文件路径
            output_path: 输出文件路径
            generated_contents: {段落索引: GeneratedContentResult列表}
            generated_images: {段落索引: 图片信息列表}，可选（已包含在result中，此处为兼容保留）

        Returns:
            输出文件路径
        """
        logger.info(
            "replace_blocks_start",
            template_path=template_path,
            output_path=output_path,
            content_count=len(generated_contents),
        )

        try:
            doc = Document(template_path)
            paragraphs_info = self.extract_paragraphs_info(template_path)

            # 按段落索引降序排序，避免插入新段落后索引错乱
            sorted_items = sorted(
                generated_contents.items(),
                key=lambda x: int(x[0]),
                reverse=True,
            )

            for para_idx_str, results in sorted_items:
                para_idx = int(para_idx_str)

                if para_idx >= len(doc.paragraphs):
                    logger.warning(
                        "paragraph_index_out_of_range",
                        index=para_idx,
                        total=len(doc.paragraphs),
                    )
                    continue

                if not results:
                    logger.warning(
                        "empty_results",
                        paragraph_index=para_idx,
                    )
                    continue

                # 判断是列表还是单一段落
                if len(results) > 1 or results[0].children:
                    # 列表段落：多个结果或有子段落
                    current_para = self._replace_with_list_results(
                        doc, para_idx, paragraphs_info, results
                    )
                else:
                    # 单一段落：直接替换
                    result = results[0]
                    paragraph = doc.paragraphs[para_idx]
                    self._replace_with_text(paragraph, result.content)
                    current_para = paragraph

                    # 插入关联的图片
                    if result.images:
                        current_para = self._insert_images_after_paragraph(
                            doc, current_para, result.images
                        )

            # 确保输出目录存在
            output_path_obj = Path(output_path)
            output_path_obj.parent.mkdir(parents=True, exist_ok=True)

            # 保存文档
            doc.save(output_path)

            logger.info(
                "replace_blocks_success",
                output_path=output_path,
            )

            return output_path

        except Exception as e:
            logger.error(
                "replace_blocks_failed",
                template_path=template_path,
                output_path=output_path,
                error=str(e),
            )
            raise

    def _replace_with_text(self, paragraph: Paragraph, content: str) -> None:
        """用文本替换段落中的内容块.

        Args:
            paragraph: 段落对象
            content: 替换内容
        """
        # 替换内容块为生成内容
        original_text = paragraph.text
        new_text = self.block_pattern.sub(content, original_text, count=1)

        # 清除段落并重新添加文本
        paragraph.clear()
        run = paragraph.add_run(new_text)

        # 保留原始字体大小，如果没有则设置默认
        if run.font.size is None:
            run.font.size = Pt(12)

    def _replace_with_list_results(
        self,
        doc: Document,
        para_idx: int,
        paragraphs_info: List[ParagraphInfo],
        results: List["GeneratedContentResult"],
    ) -> Paragraph:
        """将列表模板段落替换为生成的列表内容.

        Args:
            doc: 文档对象
            para_idx: 列表模板段落索引
            paragraphs_info: 所有段落信息
            results: 列表项生成结果列表

        Returns:
            最后插入的段落对象
        """
        if not results:
            # 空列表，清空原段落
            paragraph = doc.paragraphs[para_idx]
            paragraph.clear()
            return paragraph

        # 获取原段落样式
        original_paragraph = doc.paragraphs[para_idx]
        original_style = original_paragraph.style

        # 找到列表模板段落后的原始子段落数量（用于删除）
        original_child_count = self._count_list_children(
            paragraphs_info, para_idx
        )

        # 删除原始子段落（从后向前删除）
        for i in range(original_child_count):
            try:
                para_to_remove = doc.paragraphs[para_idx + 1]
                para_to_remove._element.getparent().remove(para_to_remove._element)
            except Exception:
                break

        # 处理第一个列表项：使用原段落
        first_item = results[0]
        original_paragraph.clear()
        run = original_paragraph.add_run(first_item.content)
        if run.font.size is None:
            run.font.size = Pt(12)

        # 为第一个列表项添加子段落
        current_paragraph = original_paragraph
        if first_item.children:
            current_paragraph = self._add_child_results(
                doc, current_paragraph, first_item.children, original_style
            )

        # 插入第一个列表项关联的图片
        if first_item.images:
            current_paragraph = self._insert_images_after_paragraph(
                doc, current_paragraph, first_item.images
            )

        # 处理剩余的列表项
        for item in results[1:]:
            # 添加列表项标题
            current_paragraph = self._insert_paragraph_after(
                doc, current_paragraph, item.content
            )
            try:
                current_paragraph.style = original_style
            except Exception:
                pass

            # 为列表项添加子段落
            if item.children:
                current_paragraph = self._add_child_results(
                    doc, current_paragraph, item.children, original_style
                )

            # 插入列表项关联的图片
            if item.images:
                current_paragraph = self._insert_images_after_paragraph(
                    doc, current_paragraph, item.images
                )

        return current_paragraph

    def _add_child_results(
        self,
        doc: Document,
        after_paragraph: Paragraph,
        children: List["GeneratedContentResult"],
        parent_style,
    ) -> Paragraph:
        """添加子段落结果.

        Args:
            doc: 文档对象
            after_paragraph: 参考段落（在此段落后插入）
            children: 子段落结果列表
            parent_style: 父段落样式

        Returns:
            最后插入的段落
        """
        current_paragraph = after_paragraph

        for child in children:
            if child.children:
                # 嵌套列表：递归处理
                current_paragraph = self._add_nested_list_results(
                    doc, current_paragraph, [child], parent_style
                )
            else:
                # 普通段落
                current_paragraph = self._insert_paragraph_after(
                    doc, current_paragraph, child.content
                )
                # 根据 style_name 设置段落样式
                try:
                    if child.style_name:
                        current_paragraph.style = doc.styles[child.style_name]
                except Exception:
                    # 样式设置失败时忽略（使用默认样式）
                    pass

                # 插入子段落关联的图片
                if child.images:
                    current_paragraph = self._insert_images_after_paragraph(
                        doc, current_paragraph, child.images
                    )

        return current_paragraph

    def _add_nested_list_results(
        self,
        doc: Document,
        after_paragraph: Paragraph,
        results: List["GeneratedContentResult"],
        parent_style,
    ) -> Paragraph:
        """添加嵌套列表内容.

        Args:
            doc: 文档对象
            after_paragraph: 参考段落（在此段落后插入）
            results: 嵌套列表结果
            parent_style: 父段落样式

        Returns:
            最后插入的段落
        """
        current_paragraph = after_paragraph

        if not results:
            return current_paragraph

        for item in results:
            # 创建列表项标题段落
            current_paragraph = self._insert_paragraph_after(
                doc, current_paragraph, item.content
            )
            # 使用 item.style_name 或 parent_style
            try:
                if hasattr(item, 'style_name') and item.style_name:
                    current_paragraph.style = doc.styles[item.style_name]
                else:
                    current_paragraph.style = parent_style
            except Exception:
                pass

            # 递归处理子段落
            if item.children:
                current_paragraph = self._add_child_results(
                    doc, current_paragraph, item.children, parent_style
                )

            # 插入列表项关联的图片
            if item.images:
                current_paragraph = self._insert_images_after_paragraph(
                    doc, current_paragraph, item.images
                )

        return current_paragraph

    def _count_list_children(
        self,
        paragraphs_info: List[ParagraphInfo],
        list_para_idx: int,
    ) -> int:
        """计算列表模板段落后的子段落数量.

        子段落是指位于列表模板之后、下一个同层级或更高层级标题之前的段落.

        Args:
            paragraphs_info: 所有段落信息
            list_para_idx: 列表模板段落索引

        Returns:
            子段落数量
        """
        count = 0
        list_para_info = None

        # 找到列表模板段落
        for info in paragraphs_info:
            if info.index == list_para_idx:
                list_para_info = info
                break

        if not list_para_info:
            return 0

        # 检查后续段落
        for info in paragraphs_info:
            if info.index <= list_para_idx:
                continue

            # 如果遇到同层级或更高层级的标题，停止计数
            if info.is_heading:
                list_level = self._get_heading_level(list_para_info.style_name)
                current_level = self._get_heading_level(info.style_name)

                if current_level <= list_level:
                    break

            count += 1

        return count

    def _get_heading_level(self, style_name: str) -> int:
        """获取标题级别.

        Args:
            style_name: 样式名称

        Returns:
            标题级别（非标题返回99）
        """
        if not style_name.startswith("Heading"):
            return 99

        try:
            level = int(style_name.replace("Heading", "").strip())
            return level
        except ValueError:
            return 99

    def _insert_paragraph_after(
        self,
        doc: Document,
        paragraph: Paragraph,
        text: str,
    ) -> Paragraph:
        """在指定段落后插入新段落.

        Args:
            doc: 文档对象
            paragraph: 参考段落
            text: 新段落文本

        Returns:
            新段落对象
        """
        from docx.oxml import OxmlElement

        # 创建新段落元素
        new_p = OxmlElement('w:p')
        paragraph._element.addnext(new_p)

        # 创建段落对象
        new_paragraph = Paragraph(new_p, paragraph._parent)

        # 添加文本
        run = new_paragraph.add_run(text)
        if run.font.size is None:
            run.font.size = Pt(12)

        return new_paragraph

    def _insert_images_after_paragraph(
        self,
        doc: Document,
        paragraph: Paragraph,
        images: List["ImageInfo"],
        max_width: float = 6.0,
    ) -> Paragraph:
        """在指定段落后插入图片.

        Args:
            doc: 文档对象
            paragraph: 参考段落（在此段落后插入）
            images: 图片信息列表（包含文件路径）
            max_width: 图片最大宽度（英寸）

        Returns:
            最后插入的段落对象
        """
        from docx.oxml import OxmlElement

        current_paragraph = paragraph

        for image_info in images:
            try:
                if not image_info.image_path:
                    logger.warning("image_path_empty", image_id=image_info.image_id)
                    continue

                image_path = Path(image_info.image_path)
                if not image_path.exists():
                    logger.warning(
                        "image_file_not_found",
                        image_id=image_info.image_id,
                        image_path=str(image_path),
                    )
                    continue

                # 创建新段落用于放置图片
                new_p = OxmlElement('w:p')
                current_paragraph._element.addnext(new_p)
                current_paragraph = Paragraph(new_p, paragraph._parent)

                # 添加图片
                run = current_paragraph.add_run()

                # 计算图片尺寸，保持宽高比
                from PIL import Image as PILImage
                with PILImage.open(image_path) as img:
                    img_width, img_height = img.size
                    aspect_ratio = img_height / img_width

                    # 如果图片宽度超过最大宽度，按比例缩小
                    if img_width > max_width * 96:  # 96 DPI
                        display_width = Inches(max_width)
                        display_height = Inches(max_width * aspect_ratio)
                    else:
                        display_width = Inches(img_width / 96)
                        display_height = Inches(img_height / 96)

                # 从文件路径添加图片
                run.add_picture(str(image_path), width=display_width, height=display_height)

                logger.info(
                    "image_inserted",
                    image_id=image_info.image_id,
                    img_method_name=image_info.method_name,
                    image_path=str(image_path),
                )

            except Exception as e:
                logger.warning(
                    "image_insert_failed",
                    image_id=image_info.image_id,
                    error=str(e),
                )

        return current_paragraph

    def get_document_info(self, doc_path: str) -> Dict[str, any]:
        """获取文档基本信息.

        Args:
            doc_path: 文档路径

        Returns:
            文档信息字典
        """
        doc = Document(doc_path)

        paragraph_count = len(doc.paragraphs)
        block_count = 0

        for paragraph in doc.paragraphs:
            matches = self.block_pattern.findall(paragraph.text)
            block_count += len(matches)

        return {
            "path": doc_path,
            "paragraph_count": paragraph_count,
            "block_count": block_count,
        }

    def validate_template(self, doc_path: str) -> Tuple[bool, Optional[str]]:
        """验证模板文件是否有效.

        Args:
            doc_path: 文档路径

        Returns:
            (是否有效, 错误信息)
        """
        try:
            path = Path(doc_path)

            if not path.exists():
                return False, f"Template file not found: {doc_path}"

            if not path.suffix.lower() in ['.docx', '.doc']:
                return False, f"Invalid file format: {path.suffix}"

            doc = Document(doc_path)

            has_blocks = False
            for paragraph in doc.paragraphs:
                if self.block_pattern.search(paragraph.text):
                    has_blocks = True
                    break

            if not has_blocks:
                logger.warning(
                    "template_no_blocks",
                    doc_path=doc_path,
                )

            return True, None

        except Exception as e:
            return False, f"Failed to validate template: {str(e)}"

    def replace_dot(self, text: str) -> str:
        extra_map = {
            '"': '"', '"': '"',
            ''': "'", ''': "'",
            '，': ','
        }
        for k, v in extra_map.items():
            text = text.replace(k, v)
        return text
