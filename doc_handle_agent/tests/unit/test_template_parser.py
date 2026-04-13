"""模板解析器单元测试."""

import json
import tempfile
from pathlib import Path

import pytest
from docx import Document

from app.core.state import ParagraphType, StaticParagraph, TemplateParagraph
from app.core.template_parser import TemplateParser, create_example_template


class TestTemplateParser:
    """模板解析器测试类."""

    @pytest.fixture
    def parser(self):
        """创建解析器实例."""
        return TemplateParser()

    @pytest.fixture
    def sample_template(self, tmp_path):
        """创建示例模板文件."""
        doc = Document()

        # 添加标题
        doc.add_heading("系统设计文档", level=1)

        # 添加模板段落
        doc.add_heading("1. 系统概述", level=2)
        doc.add_paragraph('{{"prompt":"系统的整体功能概述"}}')

        doc.add_heading("2. 功能模块", level=2)
        doc.add_paragraph('{{"prompt":"系统的主要功能模块", "list":"true", "min_length":"3", "max_length":"5"}}')

        # 添加普通段落
        doc.add_paragraph("这是一段普通文本，不包含内容块。")

        # 保存
        template_path = tmp_path / "test_template.docx"
        doc.save(template_path)

        return str(template_path)

    def test_parse_valid_template(self, parser, sample_template):
        """测试解析有效模板."""
        paragraphs = parser.parse(sample_template)

        assert len(paragraphs) == 2

        # 第一个段落（普通模板）
        assert paragraphs[0].id == "2"  # 第3个段落（索引从0开始，实际是第3个段落）
        assert paragraphs[0].is_template is True
        assert paragraphs[0].is_heading is False
        assert paragraphs[0].prompt == "系统的整体功能概述"
        assert paragraphs[0].is_list is False

        # 第二个段落（列表模板）
        assert paragraphs[1].id == "4"  # 第5个段落
        assert paragraphs[1].is_template is True
        assert paragraphs[1].is_heading is False  # 模板段落本身不是标题
        assert paragraphs[1].prompt == "系统的主要功能模块"
        assert paragraphs[1].is_list is True
        assert paragraphs[1].min_length == 3
        assert paragraphs[1].max_length == 5

    def test_parse_empty_template(self, parser, tmp_path):
        """测试解析空模板."""
        doc = Document()
        doc.add_paragraph("没有内容块的普通文档。")

        template_path = tmp_path / "empty_template.docx"
        doc.save(template_path)

        paragraphs = parser.parse(str(template_path))
        assert len(paragraphs) == 0

    def test_parse_invalid_json(self, parser, tmp_path):
        """测试解析无效JSON."""
        doc = Document()
        doc.add_paragraph('{{invalid json}}')

        template_path = tmp_path / "invalid_template.docx"
        doc.save(template_path)

        with pytest.raises(ValueError):
            parser.parse(str(template_path))

    def test_validate_template(self, parser, sample_template):
        """测试模板验证."""
        is_valid, message = parser.validate_template(sample_template)

        assert is_valid is True
        assert message is None or "Warning" in message

    def test_validate_nonexistent_file(self, parser):
        """测试验证不存在的文件."""
        is_valid, message = parser.validate_template("/nonexistent/file.docx")

        assert is_valid is False
        assert "not found" in message.lower()

    def test_preview_blocks(self, parser, sample_template):
        """预览模板段落."""
        previews = parser.preview_blocks(sample_template)

        assert len(previews) == 2
        assert previews[0]["is_template"] is True
        assert previews[0]["is_heading"] is False
        assert previews[1]["is_template"] is True
        assert previews[1]["is_list"] is True

    def test_parse_non_heading_template_no_children(self, parser, tmp_path):
        """测试解析非标题模板段落（新算法：非标题模板不收集子段落，全部为独立项）."""
        doc = Document()

        # 添加列表模板
        doc.add_heading("1. 功能模块", level=2)
        doc.add_paragraph('{{"prompt":"系统的主要功能模块", "list":"true"}}')

        # 添加子段落（混合静态和模板）
        doc.add_paragraph('    1.1 标识 {{"prompt":"模块标识"}}', style='List Paragraph')
        doc.add_paragraph('    1.2 静态描述', style='List Paragraph')
        doc.add_paragraph('    1.3 概要 {{"prompt":"功能概要"}}', style='List Paragraph')

        template_path = tmp_path / "list_template.docx"
        doc.save(template_path)

        paragraphs = parser.parse(str(template_path))

        # 新算法：非标题模板段落不收集子段落，所有模板都是独立项
        # 所以应该有 3 个模板段落（系统功能模块、模块标识、功能概要）
        assert len(paragraphs) == 3

        # 第一个段落是 "系统的主要功能模块"
        assert paragraphs[0].prompt == "系统的主要功能模块"
        assert paragraphs[0].is_list is True
        assert len(paragraphs[0].children) == 0  # 非标题模板不收集子段落

        # 第二个段落是 "模块标识"
        assert paragraphs[1].prompt == "模块标识"
        assert paragraphs[1].is_list is False

        # 第三个段落是 "功能概要"
        assert paragraphs[2].prompt == "功能概要"
        assert paragraphs[2].is_list is False

    def test_parse_heading_with_children(self, parser, tmp_path):
        """测试解析标题模板段落的子段落（新算法：只有标题模板收集子段落）."""
        doc = Document()

        # 添加标题模板（Heading 1）
        doc.add_heading('{{"prompt":"第一章 概述"}}', level=1)

        # 添加子标题模板（Heading 2）
        doc.add_heading('{{"prompt":"1.1 背景"}}', level=2)
        doc.add_paragraph('{{"prompt":"背景描述内容"}}')

        doc.add_heading('{{"prompt":"1.2 目标"}}', level=2)
        doc.add_paragraph('{{"prompt":"目标描述内容"}}')

        # 添加同级标题（Heading 1）- 这会终止前一个标题的子段落收集
        doc.add_heading('{{"prompt":"第二章 设计"}}', level=1)
        doc.add_paragraph('{{"prompt":"设计内容"}}')

        template_path = tmp_path / "heading_template.docx"
        doc.save(template_path)

        paragraphs = parser.parse(str(template_path))

        # 应该有两个顶层标题模板
        assert len(paragraphs) == 2

        # 第一个标题：第一章 概述
        assert paragraphs[0].prompt == "第一章 概述"
        assert paragraphs[0].is_heading is True
        assert len(paragraphs[0].children) == 2  # 两个子标题

        # 检查第一个标题的子段落
        first_children = paragraphs[0].children
        assert isinstance(first_children[0], TemplateParagraph)
        assert first_children[0].prompt == "1.1 背景"
        assert first_children[0].is_heading is True
        assert len(first_children[0].children) == 1  # 背景描述
        assert isinstance(first_children[0].children[0], TemplateParagraph)
        assert first_children[0].children[0].prompt == "背景描述内容"

        assert isinstance(first_children[1], TemplateParagraph)
        assert first_children[1].prompt == "1.2 目标"
        assert first_children[1].is_heading is True
        assert len(first_children[1].children) == 1  # 目标描述
        assert isinstance(first_children[1].children[0], TemplateParagraph)
        assert first_children[1].children[0].prompt == "目标描述内容"

        # 第二个标题：第二章 设计
        assert paragraphs[1].prompt == "第二章 设计"
        assert paragraphs[1].is_heading is True
        assert len(paragraphs[1].children) == 1  # 设计内容
        assert isinstance(paragraphs[1].children[0], TemplateParagraph)
        assert paragraphs[1].children[0].prompt == "设计内容"


class TestExampleTemplate:
    """示例模板测试类."""

    def test_create_example_template(self, tmp_path):
        """测试创建示例模板."""
        output_path = str(tmp_path / "example.docx")
        result_path = create_example_template(output_path)

        assert Path(result_path).exists()

        # 验证可以解析
        parser = TemplateParser()
        paragraphs = parser.parse(result_path)

        assert len(paragraphs) > 0
