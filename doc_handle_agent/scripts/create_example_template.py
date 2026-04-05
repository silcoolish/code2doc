#!/usr/bin/env python3
"""创建示例模板文件."""

from docx import Document
from docx.shared import Pt


def create_example_template(output_path: str = "templates/example_template.docx"):
    """创建示例模板文件."""
    doc = Document()

    # 标题
    doc.add_heading('系统设计文档', level=1)

    # 1. 系统概述
    doc.add_heading('1. 系统概述', level=2)
    doc.add_paragraph('{{"type":"text", "prompt":"系统的整体功能概述，包括主要业务目标和技术目标"}}')

    # 2. 系统架构
    doc.add_heading('2. 系统架构', level=2)
    doc.add_paragraph('{{"type":"text", "prompt":"系统的整体架构设计，包括技术栈、部署架构等"}}')

    # 3. 功能模块
    doc.add_heading('3. 功能模块', level=2)
    doc.add_paragraph('{{"type":"headline", "prompt":"系统的主要功能模块", "list":"true", "min_length":"3", "max_length":"10"}}')

    # 4. 核心流程
    doc.add_heading('4. 核心流程', level=2)
    doc.add_paragraph('{{"type":"text", "prompt":"系统的核心业务处理流程说明"}}')

    # 5. 数据模型
    doc.add_heading('5. 数据模型', level=2)
    doc.add_paragraph('{{"type":"text", "prompt":"系统的主要数据模型和数据库设计说明"}}')

    # 6. 接口设计
    doc.add_heading('6. 接口设计', level=2)
    doc.add_paragraph('{{"type":"headline", "prompt":"系统对外提供的主要API接口", "list":"true", "min_length":"5", "max_length":"15"}}')

    # 7. 总结
    doc.add_heading('7. 总结', level=2)
    doc.add_paragraph('{{"type":"text", "prompt":"系统的优势、特点以及后续优化方向"}}')

    # 保存
    import os
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    print(f'示例模板已创建: {output_path}')
    return output_path


if __name__ == "__main__":
    create_example_template()
